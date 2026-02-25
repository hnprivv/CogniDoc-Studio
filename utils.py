from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.output_parsers import StrOutputParser
from langchain_core.prompts import ChatPromptTemplate
from langchain_community.vectorstores import Chroma
from langchain_ollama import ChatOllama
from collections import Counter
from textblob import TextBlob
from pypdf import PdfReader
import streamlit as st
from time import time
import os
import re
import docx

class AppUtils:
    @staticmethod
    def load_css(file_path):
        """Injects custom CSS from an external file."""
        if os.path.exists(file_path):
            with open(file_path) as f:
                st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)

    @staticmethod
    def process_text(text, trim_space=False, case_style="No Change"):
        """Handles all text transformations before PDF generation."""
        if not text:
            return ""
            
        processed = text
        
        if trim_space:
            processed = " ".join(processed.split())
            
        if case_style == "UPPERCASE":
            processed = processed.upper()
        elif case_style == "lowercase":
            processed = processed.lower()
        elif case_style == "Title Case":
            processed = processed.title()
            
        return processed

    @staticmethod
    def get_sentiment(text):
        """Analyzes text to determine emotional tone."""
        if not text:
            return None
        
        blob = TextBlob(text)
        polarity = blob.sentiment.polarity
        
        # Determining label based on polarity
        if polarity > 0.1:
            label = "Positive"
            color = "green"
        elif polarity < -0.1:
            label = "Negative"
            color = "red"
        else:
            label = "Neutral"
            color = "gray"
            
        return {
            "label": label,
            "score": round(polarity, 2),
            "subjectivity": round(blob.sentiment.subjectivity, 2),
            "color": color
        }

    @staticmethod
    def get_stats(text):
        """Calculates enhanced document metrics."""
        words = text.split()
        sentiment_data = AppUtils.get_sentiment(text)
        
        return {
            "words": len(words),
            "chars": len(text),
            "reading_time": max(1, round(len(words) / 200)),
            "grade": AppUtils.calculate_readability(text, words),
            "top_terms": AppUtils.extract_keywords(words),
            "sentiment": sentiment_data
        }

    @staticmethod
    def calculate_readability(text, words):
        """Estimates Flesch-Kincaid Grade Level."""
        if not words: return "N/A"
        sentences = len(re.split(r'[.!?]+', text)) or 1
        # Simple syllable estimation (vowel groups)
        syllables = sum(len(re.findall(r'[aeiouy]+', w.lower())) for w in words)
        
        # Flesch-Kincaid Formula
        grade = 0.39 * (len(words) / sentences) + 11.8 * (syllables / len(words)) - 15.59
        return round(max(0, grade), 1)

    @staticmethod
    def extract_keywords(words):
        """Identifies the most frequent significant words."""
        # Filter common 'stop words' manually to keep it lightweight
        stops = {'the', 'and', 'is', 'in', 'it', 'of', 'to', 'for', 'with', 'a', 'an'}
        filtered = [w.lower() for w in words if w.lower().isalpha() and w.lower() not in stops]
        return Counter(filtered).most_common(3)
    
    @staticmethod
    def clear_text():
        st.session_state.pdf_generated = False
        # This explicitly clears the text_area's internal value
        if "content_area" in st.session_state:
            st.session_state["content_area"] = ""

            if "vector_db" in st.session_state:
                del st.session_state["vector_db"]
            
            if "messages" in st.session_state:
                st.session_state.messages = []

    @staticmethod
    def handle_voice_input(text_from_voice):
        """Appends voice transcription to the content area and triggers rerun."""
        if text_from_voice:
            # Check if there is existing content to append to
            current_content = st.session_state.get("content_area", "")
            if current_content:
                st.session_state["content_area"] = f"{current_content} {text_from_voice}"
            else:
                st.session_state["content_area"] = text_from_voice
            
            st.rerun()

    @staticmethod
    def handle_file_upload(placeholder):
        if st.session_state.file_import_widget is not None:
            uploaded_file = st.session_state.file_import_widget
            ext = uploaded_file.name.split('.')[-1].lower()
            extracted_text = ""

            with placeholder.container():
                bar = st.progress(0, text="Starting extraction...")
                
                try:
                    if ext == "txt":
                        extracted_text = uploaded_file.getvalue().decode("utf-8")
                        bar.progress(1.0, text="Reading text file...")

                    elif ext == "pdf":
                        reader = PdfReader(uploaded_file)
                        total = len(reader.pages)
                        for i, page in enumerate(reader.pages):
                            extracted_text += (page.extract_text() or "") + "\n"
                            bar.progress((i + 1) / total, text=f"Reading PDF page {i+1}...")
                    
                    elif ext == "docx":
                        doc = docx.Document(uploaded_file)
                        total = len(doc.paragraphs)
                        for i, para in enumerate(doc.paragraphs):
                            extracted_text += para.text + "\n"
                            if i % 5 == 0: # Update less frequently for speed
                                bar.progress((i + 1) / total, text="Reading Word document...")

                    # Update state and clean up
                    st.session_state.content_area = extracted_text
                    bar.progress(1.0, text="✅ Import Complete!")
                    time.sleep(1) 
                    placeholder.empty()
                    
                except Exception as e:
                    st.error(f"Extraction failed: {e}")

        else:
            # If the file is removed (uploader is empty), content area is cleared
            st.session_state["content_area"] = ""
    
    @staticmethod
    def parse_markdown(text):
        """
        Simple parser to handle basic Markdown bold and italic tags.
        Returns a list of tuples: (style, text_segment)
        Example: [('', 'Hello '), ('B', 'World')]
        """
        import re
        # Pattern for bold (**text**) or italic (*text*)
        pattern = r'(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)'
        parts = re.split(pattern, text)
        formatted_segments = []

        for part in parts:
            if part.startswith('***') and part.endswith('***'):
                formatted_segments.append(('BI', part.strip('*')))
            elif part.startswith('**') and part.endswith('**'):
                formatted_segments.append(('B', part.strip('*')))
            elif part.startswith('*') and part.endswith('*'):
                formatted_segments.append(('I', part.strip('*')))
            else:
                formatted_segments.append(('', part))
        return formatted_segments
    
    @staticmethod
    def get_line_type(line):
        stripped = line.strip()
        if stripped.startswith('# '):
            return 'H1', stripped[2:]
        elif stripped.startswith('## '):
            return 'H2', stripped[3:]
        elif stripped.startswith('### '): 
            return 'H3', stripped[4:]
        elif stripped.startswith('> '):    
            return 'QUOTE', stripped[2:]
        elif stripped == '---':            
            return 'HR', ''
        elif stripped.startswith('- ') or stripped.startswith('* '):
            return 'BULLET', stripped[2:]
        return 'TEXT', line
    
    @staticmethod
    def create_vector_store(text):
        """Processes text into a vector database for RAG."""
        if not text:
            return None
            
        # 1. Chunking
        splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        chunks = splitter.split_text(text)
        
        # 2. Local Embeddings
        embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        
        # 3. Create Vector Store (Ephemeral/In-Memory)
        vector_db = Chroma.from_texts(
            texts=chunks, 
            embedding=embeddings,
            collection_name=f"temp_{int(time())}"
        )
        return vector_db

    @staticmethod
    def ai_auto_format(text):
        """Uses a local AI to format text in chunks for better speed and stability."""
        if not text:
            return ""

        # 1. Initialize the local model with a capped context for speed
        llm = ChatOllama(model="llama3.2", temperature=0, num_ctx=2048, num_thread=10)
        
        # 2. Setup the Recursive Splitter
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=3000, 
            chunk_overlap=200,
            separators=["\n\n", "\n", " ", ""]
        )
        
        chunks = text_splitter.split_text(text)
        formatted_parts = []
        
        # 3. Create a progress bar in the Streamlit UI
        progress_text = "AI is formatting your document..."
        my_bar = st.progress(0, text=progress_text)

        # 4. Process each chunk individually
        for i, chunk in enumerate(chunks):
            prompt = ChatPromptTemplate.from_messages([
                ("system", "You are a professional document formatter. Transform the text into Markdown using #, ##, ###, and - bullets."),
                ("human", f"Format this section. Return ONLY the markdown:\n\n{chunk}")
            ])
            
            chain = prompt | llm | StrOutputParser()
            formatted_parts.append(chain.invoke({}))
            
            # Update progress
            percent_complete = (i + 1) / len(chunks)
            my_bar.progress(percent_complete, text=f"Processing part {i+1} of {len(chunks)}...")

        my_bar.empty()
        return "\n\n".join(formatted_parts)
    
    @staticmethod
    @st.dialog("CogniDoc", width="large")
    def show_chat_modal(*args, **kwargs):
        st.write("Ask CogniDoc AI anything about your uploaded document!")
        st.caption("Powered by Llama3.2.")
        
        # 1. Initialize message history
        if "messages" not in st.session_state:
            st.session_state.messages = []

        # 2. Create a container for the chat history
        chat_container = st.container(height=450)

        # 3. Render history inside the container
        with chat_container:
            for message in st.session_state.messages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])

        # 4. Chat Input
        if prompt := st.chat_input("What is the main summary?", key="document_chat_input"):
            # Add user message to state and render it
            st.session_state.messages.append({"role": "user", "content": prompt})
            with chat_container: # Render new message in the container
                with st.chat_message("user"):
                    st.markdown(prompt)

            # Generate AI response
            with chat_container:
                with st.chat_message("assistant"):
                    with st.spinner("Reading the document..."):
                        if st.session_state.get("vector_db"):
                            response = AppUtils.chat_with_doc(prompt, st.session_state.vector_db)
                            st.markdown(response)
                            st.session_state.messages.append({"role": "assistant", "content": response})
                        else:
                            st.error("Context not found. Please re-index.")

    @staticmethod
    def chat_with_doc(query, vector_db):
        """Standard RAG retrieval and generation logic."""
        # Find the 3 most relevant segments
        docs = vector_db.similarity_search(query, k=3)
        context = "\n\n".join([d.page_content for d in docs])
        
        llm = ChatOllama(model="llama3.2", temperature=0)
        
        # System instructions ensure it doesn't hallucinate outside the doc
        prompt = f"""
        SYSTEM: You are a strict document analysis tool. 
        1. Use ONLY the provided context to answer. 
        2. If the answer is not in the context, say "I cannot find this in the current document."
        3. Do NOT use outside knowledge or info from previous chats.

        CONTEXT:
        {context}

        QUESTION: {query}
        """
        return llm.invoke(prompt).content